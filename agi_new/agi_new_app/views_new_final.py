# Environment Creation
import json
import os
from datetime import datetime, timedelta

import chromadb
import requests
from django.conf import settings
from django.core.files.storage import default_storage
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from openai import OpenAI, OpenAIError
from rest_framework.response import Response
from rest_framework.decorators import api_view

from .database import PostgreSQLDB

db = PostgreSQLDB(dbname='test', user='test_owner', password='tcWI7unQ6REA')


# Create environment
@api_view(['POST'])
def create_environment(request):
    try:
        data = request.data
        name = data.get('name')
        api_key = data.get('api_key')
        model = data.get('model')
        temperature = data.get('temperature', 0.5)
        email = data.get('email')  # New email field

        environment_id = db.create_environment(name, api_key, model, temperature, email)
        return Response({"environment_id": environment_id}, status=201)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


# Read environment by ID
@api_view(['GET'])
def read_environment(request, environment_id):
    try:
        environment = db.read_environment(environment_id)
        if environment:
            response_data = {
                "features": [],  # Assuming no features are provided, keeping it empty
                "llm_config": {
                    "model": environment[3],  # model
                    "config": {
                        "temperature": environment[4],  # temperature
                    }
                },
                "env": {
                    "Environment_name": environment[1],  # name
                    "OPENAI_API_KEY": environment[2],  # API_KEY
                    "email": environment[5],  # New email field
                },
            }
            return Response(response_data, status=200)

        return Response({"error": "Environment not found"}, status=404)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


# Update environment by ID
@api_view(['POST'])
def update_environment(request, environment_id):
    try:
        data = request.data
        name = data.get('name')
        api_key = data.get('api_key')
        model = data.get('model')
        temperature = data.get('temperature')
        email = data.get('email')  # New email field

        updated_rows = db.update_environment(
            environment_id,
            name,
            api_key,
            model,
            temperature,
            email  # Pass email to the update function
        )

        if updated_rows:
            return Response({"message": f"Environment with ID {environment_id} updated successfully."}, status=200)
        return Response({"message": f"Environment with ID {environment_id} not updated."}, status=400)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


# Delete environment by ID
@api_view(['GET'])
def delete_environment(request, environment_id):
    try:
        deleted_rows = db.delete_environment(environment_id)
        if deleted_rows:
            return Response({"message": f"Environment with ID {environment_id} deleted successfully."}, status=200)
        return Response({"error": f"Environment with ID {environment_id} not found."}, status=404)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


# Read all environments
@api_view(['GET'])
def read_all_environments(request):
    try:
        environments = db.read_all_environments()
        if environments:
            environment_list = []
            for environment in environments:
                environment_list.append({
                    "id": environment[0],
                    "name": environment[1],
                    "api_key": environment[2],
                    "model": environment[3],
                    "temperature": environment[4],
                    "email": environment[5],  # New email field
                })
            return Response(environment_list, status=200)
        return Response({"message": "No environments found."}, status=404)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


@api_view(['POST'])
def get_all_environments_by_email(request):
    """
    Fetch all environments associated with the given email.
    """
    try:
        email = request.POST.get('email')  # Extract email from query parameters
        if not email:
            return Response({"error": "Email is required."}, status=400)

        # Fetch environments from the database
        environments = db.get_environments_by_email(email)

        # Check if environments are found
        if not environments:
            return Response({"message": "No environments found."}, status=404)

        # Structure the environments' data for JSON response
        environments_data = [
            {
                "id": env[0],
                "name": env[1],
                "api_key": env[2],
                "model": env[3],
                "temperature": env[4],
                "email": env[5]  # Include email in the response
            }
            for env in environments
        ]

        return Response({"environments": environments_data}, status=200)

    except Exception as e:
        # Log the error for further investigation
        logger.error(f"Error fetching environments: {str(e)}")
        return Response({"error": "An error occurred while fetching environments."}, status=500)


from rest_framework.decorators import api_view
from rest_framework.response import Response
import logging

# Set up logging
logger = logging.getLogger(__name__)


# Create Agent
@api_view(['POST'])
def create_agent(request):
    try:
        data = request.data
        name = data.get('name')
        system_prompt = data.get('system_prompt')
        agent_description = data.get('agent_description')
        backend_id = data.get('backend_id')  # Swapped field order: backend_id now comes before tools
        tools = data.get('tools')  # Tools used by the agent
        upload_attachment = data.get('upload_attachment', False)  # Default value set to False
        env_id = data.get('env_id')
        dynamic_agent_id = data.get('dynamic_agent_id')  # New field
        email = data.get('email')  # Added email field
        image_id = data.get('image_id')  # Added image_id field

        if not env_id:
            return Response({"error": "Environment ID is required"}, status=400)

        # Create the agent in the database
        agent_id = db.create_agent(
            name, system_prompt, agent_description, backend_id, tools, upload_attachment, env_id, dynamic_agent_id,
            email, image_id
        )

        return Response({"agent_id": agent_id}, status=201)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


# Read Agent by ID
@api_view(['GET'])
def read_agent(request, agent_id):
    try:
        agent = db.read_agent(agent_id)
        if agent:
            return Response({
                "id": agent[0],
                "name": agent[1],
                "system_prompt": agent[2],
                "agent_description": agent[3],
                "backend_id": agent[4],
                "tools": agent[5],
                "Additional_Features": {
                    "upload_attachment": agent[6],
                },
                "env_id": agent[7],
                "dynamic_agent_id": agent[8],
                "email": agent[9],
                "image_id": agent[10]  # Include image_id field
            }, status=200)
        return Response({"error": "Agent not found"}, status=404)
    except Exception as e:
        return Response({"error": str(e)}, status=400)

# Delete Agent by ID
@api_view(['GET'])
def delete_agent(request, agent_id):
    try:
        db.delete_agent(agent_id)
        return Response({"message": f"Agent with ID {agent_id} deleted successfully."}, status=204)
    except Exception as e:
        return Response({"error": str(e)}, status=400)

# Update Agent by ID
@api_view(['POST'])
def update_agent(request, agent_id):
    try:
        data = request.data
        name = data.get('name')
        system_prompt = data.get('system_prompt')
        agent_description = data.get('agent_description')
        backend_id = data.get('backend_id')
        tools = data.get('tools')
        upload_attachment = data.get('upload_attachment')
        env_id = data.get('env_id')
        dynamic_agent_id = data.get('dynamic_agent_id')
        email = data.get('email')
        image_id = data.get('image_id')  # Handle new field

        # Update agent in the database
        db.update_agent(
            agent_id, name, system_prompt, agent_description, backend_id, tools, upload_attachment, env_id,
            dynamic_agent_id, email, image_id
        )

        return Response({"message": f"Agent with ID {agent_id} updated successfully."}, status=200)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


# Read All Agents
@api_view(['GET'])
def read_all_agents(request):
    try:
        agents = db.get_all_agents()

        if not agents:
            return Response({"message": "No agents found"}, status=404)

        agents_data = [
            {
                "id": agent[0],
                "name": agent[1],
                "system_prompt": agent[2],
                "agent_description": agent[3],
                "backend_id": agent[4],
                "tools": agent[5],
                "Additional_Features": {
                    "upload_attachment": agent[6],
                },
                "env_id": agent[7],
                "dynamic_agent_id": agent[8],
                "email": agent[9],
                "image_id": agent[10]  # Include image_id field
            }
            for agent in agents
        ]

        return Response({"agents": agents_data}, status=200)
    except Exception as e:
        logger.error(f"Error fetching agents: {str(e)}")
        return Response({"error": "An error occurred while fetching agents"}, status=500)


@api_view(['POST'])
def get_all_agents_by_email(request):
    try:
        email = request.POST.get('email')
        if not email:
            return Response({"error": "Email is required."}, status=400)

        agents = db.get_agents_by_email(email)
        if not agents:
            return Response({"message": "No agents found."}, status=404)

        agents_data = [
            {
                "id": agent[0],
                "name": agent[1],
                "system_prompt": agent[2],
                "agent_description": agent[3],
                "backend_id": agent[4],
                "tools": agent[5],
                "upload_attachment": agent[6],
                "env_id": agent[7],
                "dynamic_agent_id": agent[8],
                "email": agent[9],
                "image_id": agent[10]  # Include image_id field
            }
            for agent in agents
        ]

        return Response({"agents": agents_data}, status=200)
    except Exception as e:
        logger.error(f"Error fetching agents: {str(e)}")
        return Response({"error": "An error occurred while fetching agents."}, status=500)


# ______________________________________________Functions and main api should exists between these ____________________________#
# Main api for the run environment based on the code.
from rest_framework.response import Response
from rest_framework.decorators import api_view
from rest_framework import status
import os
import base64
from django.http.response import HttpResponse
import io
import sys
import numpy as np
from plotly.graph_objs import Figure


# For travel planner agent
def extract_travel_details(prompt):
    """
    Extracts destination and number of days from a user's travel request prompt with improved accuracy.
    """
    try:
        # Normalize and clean input
        prompt = prompt.strip()

        # Improved regex patterns
        days_match = re.search(r'\bfor (\d+) (?:days|nights)\b', prompt, re.IGNORECASE)
        destination_match = re.search(r'\bto ([A-Za-z ,]+)\b', prompt, re.IGNORECASE)

        days = int(days_match.group(1)) if days_match else None
        destination = destination_match.group(1).strip() if destination_match else None

        # Further refine destination extraction
        if destination:
            destination = destination.split(',')[0].strip()  # Remove extra details after a comma

        if not destination or not days:
            return {"error": "Could not extract destination or number of days from prompt."}

        return destination, days
    except Exception as e:
        return {"error": f"Error processing prompt: {str(e)}"}


def differentiate_url(url):
    # Define specific patterns for YouTube URLs
    youtube_patterns = [
        r"^(https?://)?(www\.)?youtube\.com/watch\?v=[\w-]+(&\S*)?$",  # Matches YouTube video URLs
        r"^(https?://)?(www\.)?youtu\.be/[\w-]+$",  # Matches YouTube short URLs
        r"^(https?://)?(www\.)?youtube\.com/playlist\?list=[\w-]+(&\S*)?$",  # Matches YouTube playlist URLs
        r"^(https?://)?(www\.)?youtube\.com/channel/[\w-]+$",  # Matches YouTube channel URLs
        r"^(https?://)?(www\.)?youtube\.com/c/[\w-]+$",  # Matches YouTube custom URLs
        r"^(https?://)?(www\.)?youtube\.com/user/[\w-]+$",  # Matches YouTube user URLs
    ]
    # Check if the URL matches any YouTube pattern
    for pattern in youtube_patterns:
        if re.match(pattern, url, re.IGNORECASE):
            return "YouTube"

    # If no YouTube pattern matches, assume it's a regular website
    return "Website"


# EDUGPT Helper function.
def extract_topic_or_query(user_prompt):
    """
    Determines if the user input is a topic name or a query.
    - A topic is usually short, without question words (what, why, how).
    - A query often contains question words or special characters.
    """
    topic_pattern = r"^[A-Za-z\s]+$"  # Allows only letters and spaces (simple topic detection)
    question_words = ["what", "why", "how", "who", "when", "where", "is", "can", "does", "explain", "?"]

    user_prompt_lower = user_prompt.lower().strip()

    if any(word in user_prompt_lower for word in question_words) or not re.match(topic_pattern, user_prompt):
        return "query"  # It's a question/query
    return "topic"  # It's a valid topic name


@api_view(['POST'])
def run_openai_environment(request):
    try:
        agent_id = request.data.get('agent_id')
        user_prompt = request.data.get('prompt', '')
        file = request.FILES.get('file')
        url = request.data.get('url', '')
        file1 = request.FILES.getlist("file")

        # Retrieve agent details
        agent = db.read_agent(agent_id)

        # Retrieve the API key from the environment table using env_id
        env_details = db.read_environment(agent[7])
        openai_api_key = env_details[2]
        client = OpenAI(api_key=openai_api_key)

        result = None
        response_data = {}
        BLOG_TOOL_IDS = ['blog_post', 'audio_blog', 'video_blog', 'youtube_blog']

        # First Application: Text_to_sql
        if file and user_prompt and 'text_to_sql' in agent[4]:
            print("Text_to_Sql_condition")
            result = gen_response(file, user_prompt, client)
            # Handle the result from gen_response
            if not result:
                return Response({"error": "gen_response() returned None."},
                                status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            if "chartData" in result:
                # Visualization result
                return Response({"chartData": result["chartData"]}, status=status.HTTP_200_OK)
            elif "answer" in result:
                # Text-based result
                return Response({"answer": markdown_to_html(result["answer"])}, status=status.HTTP_200_OK)
            else:
                return Response({"error": "No valid output from gen_response."}, status=status.HTTP_400_BAD_REQUEST)

        # 3rd application:ATS Tracker
        elif file1 and user_prompt and 'ats_tracker' in agent[4]:
            print("Ats_tracker_condition")
            result = analyze_resume(user_prompt, file1, openai_api_key)
            print("result is", result)

            # Validate response format
            if "error" in result:
                # Handle error case
                return Response({"error": result["error"]}, status=status.HTTP_400_BAD_REQUEST)

            # Handle single resume analysis
            if "JD Match" in result:
                # Convert the relevant parts of the result to markdown (if needed) and return
                markdown_content = f"""
                **Job Description Match:** {result.get("JD Match", "N/A")}

                **Missing Keywords:**
                {", ".join(result.get("MissingKeywords", [])) if result.get("MissingKeywords") else "None"}

                **Profile Summary:**
                {result.get("Profile Summary", "No summary available")}

                **Suggestions:**
                {"/n".join([f"- {suggestion}" for suggestion in result.get("Suggestions", [])])}
                """
                return Response({"answers": markdown_to_html(markdown_content)}, status=status.HTTP_200_OK)

            # Handle multiple resume analysis
            elif "summary" in result and "detailed_results" in result:
                # Convert detailed results to markdown (if needed) and return
                detailed_results_html = []
                for res in result["detailed_results"]:
                    markdown_content = f"""
                    **Job Description Match:** {res.get("JD Match", "N/A")}

                    **Missing Keywords:**
                    {", ".join(res.get("MissingKeywords", [])) if res.get("MissingKeywords") else "None"}

                    **Profile Summary:**
                    {res.get("Profile Summary", "No summary available")}

                    **Suggestions:**
                    {"/n".join([f"- {suggestion}" for suggestion in res.get("Suggestions", [])])}
                    """
                    detailed_results_html.append(markdown_to_html(markdown_content))

                return Response({
                    "summary": result["summary"],
                    "detailed_results": detailed_results_html
                }, status=status.HTTP_200_OK)

            # Handle unexpected response format
            else:
                return Response({"error": "Unexpected response format from the analyzer"},
                                status=status.HTTP_500_INTERNAL_SERVER_ERROR)


        # 4th application : Chat to doc within specific page numbers and querying
        elif file and user_prompt and 'chat_to_doc_within_page_range' in agent[4]:
            print("rag condition")
            result = document_question_answering(openai_api_key, file, user_prompt)
            print("result is", result)
            if "answer" in result:
                html_response = markdown_to_html(result["answer"])
                return Response({"answer": html_response}, status=status.HTTP_200_OK)

        # 5th application: Travel planner agent
        elif user_prompt and 'travel_planner' in agent[4]:
            print("travel_planner condition")
            travel_planner_agent = TravelPlannerAgent(openai_api_key)
            destination, days = travel_planner_agent.parse_user_input(user_prompt)
            print("Destination:", destination)
            print("Days", days)
            result = generate_travel_plan(destination, days, openai_api_key)
            if not result:
                return Response({"error": "analyze_resume() returned None."},
                                status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            if "answer" in result:
                response_data["answer"] = result["answer"]
                return Response(markdown_to_html(response_data), status=status.HTTP_200_OK)

        # 6th Application:Medical Diagnosis Agent
        elif file and 'medical_diagnosis' in agent[4]:
            print("medical Diagnosis condition")
            result = run_medical_diagnosis(openai_api_key, file.read().decode("utf-8"))
            print(result)
            if not result:
                return Response({"error": "analyze_resume() returned None."},
                                status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            if "report" in result:
                return Response({"answer": markdown_to_html(result["report"])}, status=status.HTTP_200_OK)


        # 7th Application:Education Agent
        # elif user_prompt and "edu_gpt" in agent[4]:
        #     prompt_type = extract_topic_or_query(user_prompt)  # Determine if input is a topic or a query
        #
        #     # If training is in progress, block further queries
        #     if request.session.get("training", False):
        #         return Response({"message": "Training in progress. Please wait for completion."},
        #                         status=status.HTTP_200_OK)
        #
        #     # If it's a topic name, start learning
        #     if prompt_type == "topic":
        #         print("EduGPT: Learning about the topic...")
        #
        #         try:
        #             # Set training flag
        #             request.session["training"] = True
        #             request.session.modified = True
        #
        #             # Start learning
        #             result = start_learning(request, user_prompt, openai_api_key)
        #             print("Result is...................")
        #
        #             # Training completed
        #             request.session["training"] = False
        #             request.session.modified = True
        #             print(result)
        #             return Response({"answer": result["syllabus"]}, status=status.HTTP_200_OK)
        #
        #         except Exception as e:
        #             request.session["training"] = False  # Reset training flag on failure
        #             request.session.modified = True
        #             return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        #
        #     # If it's a query, proceed with chat agent
        #
        #     elif prompt_type == "query":
        #         print("EduGPT: Handling query using existing session data...")
        #         try:
        #             if "syllabus" not in request.session or "current_topic" not in request.session:
        #                 return Response({"error": "No training data found. Please start learning first."},
        #                                 status=status.HTTP_400_BAD_REQUEST)
        #
        #             # Pass API key explicitly to chat_with_agent
        #             result = chat_with_agent(request, user_prompt, openai_api_key)
        #
        #             if "error" in result:
        #                 return Response({"error": result["error"]}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        #
        #             return Response({"answer": markdown_to_html(result["assistant_response"])},
        #                             status=status.HTTP_200_OK)
        #
        #         except Exception as e:
        #             return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        #
        #     else:
        #         return Response({"error": "Invalid prompt type. Please provide a valid topic or query."},
        #                         status=status.HTTP_400_BAD_REQUEST)
        #7th application - Edugpt
        elif user_prompt and "edu_gpt" in agent[4]:
            prompt_type = extract_topic_or_query(user_prompt)  # Determine if input is a topic or a query

            # If training is in progress, block further queries
            cache_key = f"user_{request.user.id}_training" if request.user.is_authenticated else f"session_{request.session.session_key}_training"
            if cache.get(cache_key, False):
                return Response({"message": "Training in progress. Please wait for completion."},
                                status=status.HTTP_200_OK)

            # If it's a topic name, start learning
            if prompt_type == "topic":
                print("EduGPT: Learning about the topic...")

                try:
                    # Set training flag in cache
                    cache.set(cache_key, True, timeout=3600)

                    # Start learning
                    result = start_learning(request, user_prompt, openai_api_key)
                    print("Result is...................")

                    # Training completed
                    cache.set(cache_key, False, timeout=3600)
                    print(result)
                    return Response({"answer": result["syllabus"]}, status=status.HTTP_200_OK)

                except Exception as e:
                    cache.set(cache_key, False, timeout=3600)  # Reset training flag on failure
                    return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            # If it's a query, proceed with chat agent
            elif prompt_type == "query":
                print("EduGPT: Handling query using existing cache data...")
                try:
                    cache_key = f"user_{request.user.id}_edu_data" if request.user.is_authenticated else f"session_{request.session.session_key}_edu_data"
                    edu_data = cache.get(cache_key)

                    if not edu_data or "syllabus" not in edu_data or "current_topic" not in edu_data:
                        return Response({"error": "No training data found. Please start learning first."},
                                        status=status.HTTP_400_BAD_REQUEST)

                    # Pass API key explicitly to chat_with_agent
                    result = chat_with_agent(request, user_prompt, openai_api_key)

                    if "error" in result:
                        return Response({"error": result["error"]}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

                    return Response({"answer": markdown_to_html(result["assistant_response"])},
                                    status=status.HTTP_200_OK)

                except Exception as e:
                    return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            else:
                return Response({"error": "Invalid prompt type. Please provide a valid topic or query."},
                                status=status.HTTP_400_BAD_REQUEST)


        # 8th Application:medical Image processing
        elif file and 'image_processing' in agent[4]:
            print("image_processing_condition")
            result = medical_image_analysis(openai_api_key, file)

            if not result:
                return Response({"error": "analyze_resume() returned None."},
                                status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            if "result" in result:
                # response_data["result"] = result["result"]
                # response_data["simplified_explanation"] = result["simplified_explanation"]
                return Response({"answer": markdown_to_html(result["result"])}, status=status.HTTP_200_OK)

        # 9th Application:
        elif file and user_prompt and 'image_answering' in agent[4]:
            print("image_answering_condition")
            result = visual_question_answering(openai_api_key, file, user_prompt)
            if not result:
                return Response({"error": "analyze_resume() returned None."},
                                status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            if "answer" in result:
                return Response({"answer": markdown_to_html(result["answer"])}, status=status.HTTP_200_OK)


        # 10th application-Blog agents:
        elif url and user_prompt:
            print("Getting the url.............")
            url_type = differentiate_url(url)
            print(url_type)
            # Handling based on URL type
            if url_type == "YouTube":
                if any(tool_id in agent[4] for tool_id in BLOG_TOOL_IDS):
                    result = generate_blog_from_yt_url(user_prompt, url, 'blog_post', openai_api_key)
                elif 'linkedin_post' in agent[4]:
                    result = generate_blog_from_yt_url(user_prompt, url, 'linkedin_post', openai_api_key)
            else:  # General Website URL
                if any(tool_id in agent[4] for tool_id in BLOG_TOOL_IDS):
                    result = generate_blog_from_url(user_prompt, url, 'blog_post', openai_api_key)
                    print(result)
                elif 'linkedin_post' in agent[4]:
                    result = generate_blog_from_url(user_prompt, url, 'linkedin_post', openai_api_key)

            if isinstance(result, dict):
                response_data = markdown_to_html(result["content"])
                return Response({"answer": response_data}, status=status.HTTP_200_OK)


        # Blog Generation through files
        elif file and user_prompt:
            if any(tool_id in agent[4] for tool_id in BLOG_TOOL_IDS):
                print("function calling here")
                result = generate_blog_from_file(user_prompt, file, 'blog_post', openai_api_key)
                print(result)
            elif 'linkedin_post' in agent[4]:
                result = generate_blog_from_file(user_prompt, file, 'linkedin_post', openai_api_key)

            if isinstance(result, dict):
                response_data = markdown_to_html(result["content"])
                return Response({"answer": response_data}, status=status.HTTP_200_OK)


        # Second Application: Synthetic_data_generator
        elif 'synthetic_data_generation' in agent[4]:
            print("Synthetic_data_condition")
            if user_prompt and file:
                print("Extend data condition")
                result = handle_synthetic_data_from_excel(file, openai_api_key, user_prompt)
            elif file:
                print("missing_data_condition")
                result = handle_fill_missing_data(file, openai_api_key)
            else:
                print("New data condition")
                result = handle_synthetic_data_for_new_data(user_prompt, openai_api_key)

            if result and "data" in result:
                return Response(result, status=status.HTTP_200_OK)
            elif result and "error" in result:
                return Response(result, status=status.HTTP_400_BAD_REQUEST)
            else:
                return Response({"error": "An unexpected error occurred."},
                                status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        else:
            return Response({"error": "No valid tool found for the given input."}, status=status.HTTP_400_BAD_REQUEST)

    except Exception as e:
        return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


def markdown_to_html(md_text):
    html_text = markdown.markdown(md_text)
    return html_text


##1st Application:Text-to-sql(Both text and graph)
USER = 'test_owner'
PASSWORD = 'tcWI7unQ6REA'
HOST = 'ep-yellow-recipe-a5fny139.us-east-2.aws.neon.tech:5432'
DATABASE = 'test'


@csrf_exempt
def gen_response(file, query, client):
    table_name = file.name.split('.')[0]
    print(table_name)

    # Define upload directory
    upload_dir = os.path.join(settings.MEDIA_ROOT, 'upload')
    print(upload_dir)
    os.makedirs(upload_dir, exist_ok=True)

    # File path within upload directory
    file_path = os.path.join(upload_dir, file.name)
    with default_storage.open(file_path, 'wb+') as f:
        for chunk in file.chunks():
            f.write(chunk)

    df = file_to_sql(file_path, table_name, USER, PASSWORD, HOST, DATABASE)
    print(df.head(5))

    # Generate CSV metadata
    csv_metadata = {"columns": df.columns.tolist()}
    metadata_str = ", ".join(csv_metadata["columns"])
    if not query:
        return JsonResponse({"error": "No query provided"}, status=400)

    graph_keywords = [
        "plot", "graph", "visualize", "visualization", "scatter", "bar chart",
        "line chart", "histogram", "pie chart", "bubble chart", "heatmap", "box plot",
        "generate chart", "create graph", "draw", "trend", "correlation"
    ]

    # Decide whether the query is text-based or graph-based
    if any(keyword in query.lower() for keyword in graph_keywords):
        # Graph-related prompt
        print("if_condition------------------")
        print(query)
        prompt_eng = (
            f"You are an AI specialized in data analytics and visualization."
            f"Data used for analysis is stored in a pandas DataFrame named `df`. "
            f"The DataFrame `df` contains the following attributes: {metadata_str}. "
            f"Based on the user's query, generate Python code using Plotly to create the requested type of graph "
            f"(e.g., bar, pie, scatter, etc.) using the data in the DataFrame `df`. "
            f"The graph must utilize the data within `df` as appropriate for the query. "
            f"If the user does not specify a graph type, decide whether to generate a line or bar graph based on the situation."
            f"Every graph must include a title, axis labels (if applicable), and appropriate colors for better visualization."
            f"The graph must have a white background for both the plot and paper. "
            f"The code must output a Plotly 'Figure' object stored in a variable named 'fig', "
            f"and include 'data' and 'layout' dictionaries compatible with React. "
            f"The user asks: {query}."
        )

        # Call AI to generate the code
        chat = generate_code(prompt_eng, client)
        print("Generated code from AI:")
        print(chat)

        # Check for valid Plotly code in the AI response
        if 'import' in chat:
            namespace = {'df': df}  # Pass `df` into the namespace
            try:
                # Execute the generated code
                exec(chat, namespace)

                # Retrieve the Plotly figure from the namespace
                fig = namespace.get("fig")

                if fig and isinstance(fig, Figure):
                    # Convert the Plotly figure to JSON
                    chart_data = fig.to_plotly_json()

                    # Ensure JSON serialization by converting NumPy arrays to lists
                    def make_serializable(obj):
                        if isinstance(obj, np.ndarray):
                            return obj.tolist()
                        elif isinstance(obj, dict):
                            return {k: make_serializable(v) for k, v in obj.items()}
                        elif isinstance(obj, list):
                            return [make_serializable(v) for v in obj]
                        return obj

                    # Recursively process the chart_data
                    chart_data_serializable = make_serializable(chart_data)

                    # Return the structured response to the frontend
                    return {"chartData": chart_data_serializable}
                else:
                    return {"message": "No valid Plotly figure found."}
            except Exception as e:
                return {"message": f"Error executing code: {str(e)}"}
        else:
            return {"message": "AI response does not contain valid code."}


    else:
        # Text-based prompt
        print("else_condition------------------")
        print(query)
        prompt_eng = (
            f"""
                You are a Python expert focused on answering user queries about data preprocessing and analysis. Always strictly adhere to the following rules:

                1. Data-Driven Queries:
                    If the user's query is related to data processing or analysis, assume the `df` DataFrame in memory contains the actual uploaded data from the file "{file.name}" with the following columns: {metadata_str}.

                    For such queries:
                    - Generate Python code that directly interacts with the `df` DataFrame to provide accurate results strictly based on the data in the dataset.
                    - Do not make any assumptions or provide any example outputs.
                    - Ensure all answers are derived from actual calculations on the `df` DataFrame.
                    - Include concise comments explaining key steps in the code.
                    - Exclude any visualization, plotting, or assumptions about the data.

                    Example:

                    Query: "How many rows have 'Column1' > 100?"
                    Response:
                    ```python
                    # Count rows where 'Column1' > 100
                    count_rows = df[df['Column1'] > 100].shape[0]

                    # Output the result
                    print(count_rows)
                    ```

                2. Invalid or Non-Data Queries:
                    If the user's query is unrelated to data processing or analysis, or it cannot be answered using the dataset, respond with an appropriate print statement indicating the limitation. For example:

                    Query: "What is AI?"
                    Response:
                    ```python
                    print("This question is unrelated to the uploaded data. Please ask a data-specific query.")
                    ```

                3. Theoretical Concepts:
                    If the user asks about theoretical concepts in data science or preprocessing (e.g., normalization, standardization), respond with a concise explanation. Keep the response focused and accurate.

                    Example:

                    Query: "What is normalization in data preprocessing?"
                    Response:
                    ```python
                    print("Normalization is a data preprocessing technique used to scale numeric data within a specific range, typically [0, 1], to ensure all features contribute equally to the model.")
                    ```

                User query: {query}.
            """
        )
        # Generate text-related code
        code = generate_code(prompt_eng, client)
        print("Generated code from AI (Text):")
        print(code)

        # Execute the generated code with the dataset
        result = execute_py_code(code, df)
        return {"answer": result}


def file_to_sql(file_path, table_name, user, password, host, db_name):
    import pandas as pd
    import os
    from sqlalchemy import create_engine

    # engine = create_engine(f"postgresql://{user}:{password}@{host}/{db_name}")
    engine = create_mysql_engine(user, password, host, db_name)

    if not table_name:
        table_name = os.path.splitext(os.path.basename(file_path))[0]

    file_extension = os.path.splitext(file_path)[-1].lower()
    if file_extension == '.xlsx':
        df = pd.read_excel(file_path)
    elif file_extension == '.csv':
        df = pd.read_csv(file_path)
    else:
        raise ValueError("Unsupported file format. Please provide an Excel (.xlsx) or CSV (.csv) file.")

    df.to_sql(table_name, con=engine, if_exists='replace', index=False)
    return df


def create_mysql_engine(user, password, host, db_name):
    from sqlalchemy import create_engine, text

    if db_name:
        connection_str = f'postgresql://{user}:{password}@{host}/{db_name}'
    else:
        connection_str = f'postgresql://{user}:{password}@{host}/'
    engine = create_engine(connection_str)
    return engine


# Function to generate code from OpenAI API
def generate_code(prompt_eng, client):
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt_eng}
        ]
    )
    all_text = ""
    for choice in response.choices:
        message = choice.message
        chunk_message = message.content if message else ''
        all_text += chunk_message
    print(all_text)
    if "```python" in all_text:
        code_start = all_text.find("```python") + 9
        code_end = all_text.find("```", code_start)
        code = all_text[code_start:code_end]
    else:
        code = all_text
    return code


def execute_py_code(code, df):
    # Create a string buffer to capture the output
    buffer = io.StringIO()
    sys.stdout = buffer

    # Create a local namespace for execution
    local_vars = {'df': df}

    try:
        # Execute the code
        exec(code, globals(), local_vars)

        # Get the captured output
        output = buffer.getvalue().strip()

        # If there's no output, try to get the last evaluated expression
        if not output:
            last_line = code.strip().split('\n')[-1]
            if not last_line.startswith(('print', 'return')):
                output = eval(last_line, globals(), local_vars)
                print(output)
    except Exception as e:
        output = f"Error executing code: {str(e)}"
    finally:
        # Reset stdout
        sys.stdout = sys.__stdout__

    return str(output)


##1st application-Text to sql

# from wyge.prebuilt_agents.sql import DataQueryAgent
# from django.conf import settings
#
#
# def sql_query_view(api_key,uploaded_file,query_text):
#     context = {"tables": [], "query_history": []}
#     agent = DataQueryAgent(openai_api_key=api_key, model_name="gpt-4o-mini")
#
#     # Handle file upload
#     if uploaded_file:
#         file_path = default_storage.save(f"uploads/{uploaded_file.name}", ContentFile(uploaded_file.read()))
#         abs_file_path = os.path.join(settings.MEDIA_ROOT, file_path)
#
#         success = agent.load_data(file_path=abs_file_path)
#         if success:
#             context["tables"].append(uploaded_file.name)  # Mock table tracking
#             context["message"] = "Data loaded successfully!"
#         else:
#             context["error"] = "Failed to load data."
#
#         os.remove(abs_file_path)  # Clean up uploaded file
#
#     # Handle Query Execution
#     if query_text:
#         try:
#             result = agent.query(query_text)
#             context["query_history"].append({"question": result['input'], "result": result['output']})
#             context["query_result"] = result['output']
#             return context["query_result"]
#         except Exception as e:
#             context["error"] = f"Error processing query: {str(e)}"


## 2nd Application(Synthetic-data-generator) ---Create and extend data
import pandas as pd
import tempfile
import re
from wyge.prebuilt_agents.synthetic_data_generator import generate_synthetic_data, generate_data_from_text, \
    fill_missing_data_in_chunk


# 1.New data generation
def handle_synthetic_data_for_new_data(user_prompt, openai_api_key):
    try:
        # Extract number of rows from the prompt
        num_rows = extract_num_rows_from_prompt(user_prompt)
        print("number of rows is..................")
        print(num_rows)
        if num_rows is None:
            return JsonResponse({"error": "Number of rows or records not found in the prompt."}, status=400)

        # Extract column names from the prompt
        column_names = extract_columns_from_prompt(user_prompt)
        print(column_names)
        if not column_names:
            return JsonResponse({"error": "No field names found in the prompt."}, status=400)

        # Generate synthetic data using the column names and the number of rows
        generated_df = generate_data_from_text(openai_api_key, user_prompt, column_names, num_rows=num_rows)

        # Convert the generated DataFrame to CSV format
        combined_csv = generated_df.to_csv(index=False)

        return {
            "data": combined_csv
        }

    except Exception as e:
        return {"error": str(e)}


def extract_num_rows_from_prompt(user_prompt):
    match = re.search(r'(\d+)\s+(rows|records)', user_prompt, re.IGNORECASE)
    if match:
        return int(match.group(1))
    return None


def extract_columns_from_prompt(user_prompt):
    # Look for all possible field identifier formats followed by the column names
    match = re.search(r'(field names|column names|fields|columns|field_names|column_names):\s*([a-zA-Z0-9_,\s\.]+)',
                      user_prompt, re.IGNORECASE)

    if match:
        # Extract the part containing column names
        raw_columns = match.group(2).split(',')
    else:
        return []

    # Format each column name (remove spaces, convert to snake_case, lowercase)
    formatted_columns = [
        re.sub(r'[^a-zA-Z0-9]', '_', col.strip()).lower()
        for col in raw_columns
    ]

    # Remove empty column names and ensure no duplicates
    formatted_columns = list(filter(bool, formatted_columns))
    return list(dict.fromkeys(formatted_columns))  # Remove duplicates


def ensure_upload_folder():
    upload_folder = "uploads"
    if not os.path.exists(upload_folder):
        os.makedirs(upload_folder)
    return upload_folder


# Extended data
def handle_synthetic_data_from_excel(file, openai_api_key, user_prompt):
    try:
        file_extension = os.path.splitext(file.name)[1].lower()
        upload_folder = ensure_upload_folder()
        file_path = os.path.join(upload_folder, file.name)

        if file_extension == ".xlsx":
            original_df = pd.read_excel(file)
        elif file_extension == ".csv":
            original_df = pd.read_csv(file)
        else:
            return JsonResponse({"error": "Unsupported file format. Please upload an Excel or CSV file."}, status=400)

        # Save file to uploads folder
        if file_extension == ".xlsx":
            original_df.to_excel(file_path, index=False)
        elif file_extension == ".csv":
            original_df.to_csv(file_path, index=False)

        num_rows = extract_num_rows_from_prompt(user_prompt)
        generated_df = generate_synthetic_data(openai_api_key, file_path, num_rows)

        combined_df = pd.concat([original_df, generated_df], ignore_index=True)
        combined_csv = combined_df.to_csv(index=False)

        return {"data": combined_csv}
    except Exception as e:
        return {"error": str(e)}


# Missing Data
def handle_fill_missing_data(file, openai_api_key):
    try:
        file_extension = os.path.splitext(file.name)[1].lower()
        upload_folder = ensure_upload_folder()
        file_path = os.path.join(upload_folder, file.name)

        if file_extension == ".xlsx":
            original_df = pd.read_excel(file)
        elif file_extension == ".csv":
            original_df = pd.read_csv(file)
        else:
            return JsonResponse({"error": "Unsupported file format. Please upload an Excel or CSV file."}, status=400)

        # Save file to uploads folder
        if file_extension == ".xlsx":
            original_df.to_excel(file_path, index=False)
        elif file_extension == ".csv":
            original_df.to_csv(file_path, index=False)

        filled_df = fill_missing_data_in_chunk(openai_api_key, file_path)
        combined_csv = filled_df.to_csv(index=False)

        return {"data": combined_csv}
    except Exception as e:
        return {"error": str(e)}


# 3rd application-ATS Tracker##########
# from wyge.prebuilt_agents.resume_analyser import ResumeAnalyzer
# def analyze_resume(job_description, resume_file, api_key):
#     try:
#         if not job_description or not resume_file:
#             return {"error": "Missing job description or resume file"}
#
#         # Initialize the Resume Analyzer
#         analyzer = ResumeAnalyzer(api_key)
#
#         # Extract text from the uploaded resume
#         resume_text = analyzer.extract_text_from_pdf(resume_file)
#
#         # Analyze the resume against the job description
#         result = analyzer.analyze_resume(resume_text, job_description)
#
#         return {
#             "answer": result
#         }
#     except Exception as e:
#         return {"error": str(e)}

# Multiple Resume Tracking
from wyge.prebuilt_agents.resume_analyser import ResumeAnalyzer


def analyze_resume(job_description, resume_files, api_key):
    try:
        if not job_description or not resume_files:
            return {"error": "Missing job description or resume files"}

        # Initialize the Resume Analyzer
        analyzer = ResumeAnalyzer(api_key)

        # Check if it's a single resume or multiple resumes
        if isinstance(resume_files, list):
            # Multiple resume analysis
            results = []
            for resume_file in resume_files:
                # Extract text from the uploaded resume
                resume_text = analyzer.extract_text_from_pdf(resume_file)

                # Analyze the resume against the job description
                result = analyzer.analyze_resume(resume_text, job_description)

                # Add filename to the result for identification
                result["filename"] = resume_file.name
                results.append(result)

            # Create a summary table
            summary_data = []
            for i, result in enumerate(results):
                summary_data.append({
                    "Rank": i + 1,
                    "Resume": result.get("filename", "Unknown"),
                    "Match %": result.get("JD Match", "0%"),
                    "Missing Keywords": ", ".join(result.get("MissingKeywords", [])) if result.get(
                        "MissingKeywords") else "None"
                })

            return {
                "summary": summary_data,
                "detailed_results": results
            }
        else:
            # Single resume analysis
            resume_text = analyzer.extract_text_from_pdf(resume_files)
            result = analyzer.analyze_resume(resume_text, job_description)

            return {
                "JD Match": result.get("JD Match", "0%"),
                "Missing Keywords": result.get("MissingKeywords", []),
                "Profile Summary": result.get("Profile Summary", "No summary available"),
                "Suggestions": result.get("Suggestions", [])
            }

    except Exception as e:
        return {"error": str(e)}


# Application 4: Chat to doc (rag) with page extraction from start to end
from wyge.prebuilt_agents.rag import RAGApplication
import tempfile


def document_question_answering(api_key, uploaded_file, query):
    try:
        if not api_key or not uploaded_file:
            return {"error": "Missing API key or document file"}

        # Ensure the uploads folder exists
        upload_folder = ensure_upload_folder()

        # Save the uploaded file to the uploads folder
        file_path = os.path.join(upload_folder, uploaded_file.name)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.read())

        # Parse page range
        page_range = "0-25"
        parsed_page_range = None
        if page_range:
            try:
                start, end = map(int, page_range.split('-'))
                parsed_page_range = (start, end)
            except ValueError:
                return {"error": "Invalid page range format. Use 'start-end'"}

        # Initialize RAG application
        rag_app = RAGApplication(file_paths=[file_path], openai_api_key=api_key, page_range=parsed_page_range)
        response = rag_app.query(query)
        # Debug: Check the type of response
        print("Response from RAGApplication:", response)
        print("Response type:", type(response))

        # Ensure response is a string
        if isinstance(response, dict):
            if "answer" in response:
                response = response["answer"]
            else:
                return {"error": "Unexpected response format from RAGApplication: missing 'answer' key"}
        elif not isinstance(response, str):
            response = str(response)  # Convert to string if not already

        return {"answer": response}

    except Exception as e:
        return {"error": f"Error occurred: {str(e)}"}


# 5th application : Travel planner agent
from wyge.prebuilt_agents.travel_planner import TravelPlannerAgent


def generate_travel_plan(destination, days, openai_api_key):
    try:
        if not destination:
            return {"error": "Missing destination"}
        if not days:
            return {"error": "Days required"}
        print("In function destination", destination)
        print("In function_days", days)

        # Get API keys with fallback values
        # weather_api_key =weather_api_key  # get_api_key("WEATHER_API_KEY", "f701a9b1299fa3bbb07471570c730090")
        # geolocation_api_key = geolocation_api_key  #get_api_key("GEOCODING_API_KEY", "2f85379af3084ad1a9fc724dfa71b041")
        openai_api_key = openai_api_key

        # Initialize travel planner agent
        agent = TravelPlannerAgent(
            openai_api_key=openai_api_key
        )

        # Generate travel plan
        prompt = f"I need a travel plan to {destination} for {days} days"
        travel_plan = agent.generate_travel_plan(prompt)

        return {"answer": travel_plan}

    except Exception as e:
        return {"error": f"Error occurred: {str(e)}"}


## 6th application:Medical diagnosis agent
from wyge.prebuilt_agents.medical_diagnosis import Cardiologist, Psychologist, Pulmonologist


def run_medical_diagnosis(api_key, medical_report):
    try:
        specialists = {
            "Cardiologist": Cardiologist(medical_report, api_key),
            "Psychologist": Psychologist(medical_report, api_key),
            "Pulmonologist": Pulmonologist(medical_report, api_key)
        }

        diagnoses = {}
        report_content = "Medical Diagnosis Report\n===========================\n"

        for specialist, agent in specialists.items():
            diagnosis = agent.run()
            diagnoses[specialist] = diagnosis
            report_content += f"Specialist: {specialist}\nDiagnosis:\n{diagnosis}\n\n"

        return {
            "diagnosis": diagnoses,
            "report": report_content
        }
    except Exception as e:
        return {"error": str(e)}


## 7th application: Education Agent
from langchain.chat_models import ChatOpenAI
from wyge.prebuilt_agents.teaching_agent import teaching_agent_fun
from wyge.prebuilt_agents.generating_syllabus import generate_syllabus
from django.http import JsonResponse
from rest_framework.response import Response
from rest_framework import status

from django.core.cache import cache


def start_learning(request, topic, api_key):
    """
    Starts the learning process by generating a syllabus and initializing cache data.
    """
    try:
        if not topic:
            return {"error": "Topic is required."}

        # Generate the syllabus
        syllabus = generate_syllabus(api_key, topic, "Focus on providing a clear learning path.")

        # Extract only the content if it's an AIMessage
        if hasattr(syllabus, "content"):
            syllabus = syllabus.content  # ✅ Extract text only

        # Store syllabus and topic in cache
        cache_key = f"user_{request.user.id}_edu_data" if request.user.is_authenticated else f"session_{request.session.session_key}_edu_data"
        cache.set(cache_key, {
            "syllabus": syllabus,
            "current_topic": topic,
            "messages": []  # Initialize conversation history
        }, timeout=3600)  # Set a timeout for the cache (e.g., 1 hour)

        return {"topic": topic, "syllabus": syllabus}

    except Exception as e:
        return {"error": str(e)}


# def start_learning(request, topic, api_key):
#     """
#     Starts the learning process by generating a syllabus and initializing session data.
#     """
#     try:
#         if not topic:
#             return {"error": "Topic is required."}
#
#         # Generate the syllabus
#         syllabus = generate_syllabus(api_key, topic, "Focus on providing a clear learning path.")
#
#         # Extract only the content if it's an AIMessage
#         if hasattr(syllabus, "content"):
#             syllabus = syllabus.content  # ✅ Extract text only
#
#         # Store syllabus and topic in session (JSON serializable)
#         request.session["syllabus"] = syllabus
#         request.session["current_topic"] = topic
#         request.session["messages"] = []  # Initialize conversation history
#         request.session.modified = True  # Ensure session update
#
#         return {"topic": topic, "syllabus": syllabus}
#
#     except Exception as e:
#         return {"error": str(e)}


# def chat_with_agent(request, user_input, api_key):
#     """
#     Handles user interaction with the teaching agent, ensuring JSON-safe session storage.
#     """
#     try:
#         if "syllabus" not in request.session or "current_topic" not in request.session:
#             return {"error": "No training data found. Please start learning first."}
#
#         # Recreate the teaching agent dynamically
#         teaching_agent = teaching_agent_fun(api_key)
#
#         # Retrieve messages from the session
#         conversation_history = request.session.get("messages", [])
#
#         # Rebuild past conversation
#         for msg in conversation_history:
#             if msg["role"] == "user":
#                 teaching_agent["add_user_message"](msg["content"])
#             elif msg["role"] == "assistant":
#                 teaching_agent["add_ai_message"](msg["content"])
#
#         # Add new user message
#         teaching_agent["add_user_message"](user_input)
#
#         # Generate AI response
#         response = teaching_agent["generate_response"]()
#
#         # Extract AI response content
#         response_content = response.content if hasattr(response, "content") else str(response)
#
#         # Update conversation history with JSON-safe data
#         conversation_history.append({"role": "user", "content": user_input})
#         conversation_history.append({"role": "assistant", "content": response_content})
#
#         # Store updated conversation in session
#         request.session["messages"] = conversation_history
#         request.session.modified = True  # Ensure session update
#
#         return {"user_message": user_input, "assistant_response": response_content}
#
#     except Exception as e:
#         return {"error": str(e)}

def chat_with_agent(request, user_input, api_key):
    """
    Handles user interaction with the teaching agent, ensuring cache-based session storage.
    """
    try:
        cache_key = f"user_{request.user.id}_edu_data" if request.user.is_authenticated else f"session_{request.session.session_key}_edu_data"
        edu_data = cache.get(cache_key)

        if not edu_data or "syllabus" not in edu_data or "current_topic" not in edu_data:
            return {"error": "No training data found. Please start learning first."}

        # Recreate the teaching agent dynamically
        teaching_agent = teaching_agent_fun(api_key)

        # Retrieve messages from the cache
        conversation_history = edu_data.get("messages", [])

        # Rebuild past conversation
        for msg in conversation_history:
            if msg["role"] == "user":
                teaching_agent["add_user_message"](msg["content"])
            elif msg["role"] == "assistant":
                teaching_agent["add_ai_message"](msg["content"])

        # Add new user message
        teaching_agent["add_user_message"](user_input)

        # Generate AI response
        response = teaching_agent["generate_response"]()

        # Extract AI response content
        response_content = response.content if hasattr(response, "content") else str(response)

        # Update conversation history with JSON-safe data
        conversation_history.append({"role": "user", "content": user_input})
        conversation_history.append({"role": "assistant", "content": response_content})

        # Store updated conversation in cache
        edu_data["messages"] = conversation_history
        cache.set(cache_key, edu_data, timeout=3600)  # Update cache with new data

        return {"user_message": user_input, "assistant_response": response_content}

    except Exception as e:
        return {"error": str(e)}


# 8th application:medical image proccessing agent
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from wyge.prebuilt_agents.medical_image_processing import MedicalImageAnalyzer
import os


def medical_image_analysis(api_key, uploaded_file):
    # Ensure the file is in an allowed format
    allowed_extensions = {".jpg", ".jpeg", ".png"}
    file_extension = os.path.splitext(uploaded_file.name)[1].lower()

    if file_extension not in allowed_extensions:
        return {"error": "Invalid file format. Please upload a JPG, JPEG, or PNG file."}

    # Save file temporarily
    file_path = default_storage.save(
        "temp/" + uploaded_file.name, ContentFile(uploaded_file.read())
    )
    full_path = os.path.join(default_storage.location, file_path)
    print(full_path)

    # Analyze the image
    analyzer = MedicalImageAnalyzer(api_key=api_key)
    result = analyzer.analyze_image(full_path)
    simplified_explanation = analyzer.simplify_explanation(result)

    # Cleanup temporary file
    default_storage.delete(file_path)

    return {"result": result, "simplified_explanation": simplified_explanation}


# 9th Application:Image answering application
from wyge.prebuilt_agents.vqa import VisualQA
import os


def visual_question_answering(api_key, uploaded_file, question):
    # Ensure the file is in an allowed format
    allowed_extensions = {".jpg", ".jpeg", ".png"}
    file_extension = os.path.splitext(uploaded_file.name)[1].lower()

    if file_extension not in allowed_extensions:
        return {"error": "Invalid file format. Please upload a JPG, JPEG, or PNG file."}

    # Save file temporarily
    file_path = default_storage.save(
        "temp/" + uploaded_file.name, ContentFile(uploaded_file.read())
    )
    full_path = os.path.join(default_storage.location, file_path)

    # Initialize VQA agent
    vqa = VisualQA(api_key=api_key)

    # Get answer from the model
    response = vqa.ask(image_path=full_path, question=question)

    # Cleanup temporary file
    default_storage.delete(file_path)

    print(response)
    return {"answer": response}


# 10th Application:Blog generation with website,youtube and Audio
from wyge.prebuilt_agents.social_media_agents import ResearchAgent, BlogAgent, LinkedInAgent
import os


def generate_blog_from_url(prompt, url, option, api_key):
    if option == 'blog_post':
        if not url:
            return {'error': 'Website URL is required'}
        research_agent = ResearchAgent(api_key=api_key)
        blog_agent = BlogAgent(api_key=api_key)

        content = research_agent.research_website(prompt, url)
        blog_post = blog_agent.generate_blog(prompt, str(content))
        return {'content': blog_post}

    elif option == 'linkedin_post':
        if not url:
            return {'error': 'Website URL is required'}
        research_agent = ResearchAgent(api_key=api_key)
        linkedin_agent = LinkedInAgent(api_key=api_key)

        content = research_agent.research_website(prompt, url)
        linkedin_post = linkedin_agent.generate_linkedin_post(prompt, str(content))

        return {'content': linkedin_post}


# Generate content from URL (for blog or LinkedIn post)
def generate_blog_from_yt_url(prompt, url, option, api_key):
    if option == 'blog_post':
        if not url:
            return {'error': 'YouTube video URL is required'}
        research_agent = ResearchAgent(api_key=api_key)
        blog_agent = BlogAgent(api_key=api_key)

        content = research_agent.extract_transcript_from_yt_video(url)
        print("content.............", content)
        blog_post = blog_agent.generate_blog(prompt, str(content))

        return {'content': blog_post}

    elif option == 'linkedin_post':
        if not url:
            return {'error': 'YouTube video URL is required'}
        research_agent = ResearchAgent(api_key=api_key)
        linkedin_agent = LinkedInAgent(api_key=api_key)

        content = research_agent.extract_transcript_from_yt_video(url)
        linkedin_post = linkedin_agent.generate_linkedin_post(prompt, str(content))

        return {'content': linkedin_post}


# Generate content from file (for blog or LinkedIn post)
def generate_blog_from_file(prompt, file, option, api_key):
    if option == 'blog_post':
        if not file:
            return {'error': 'Audio file is required'}

        temp_path = f"temp_{file.name}"
        with open(temp_path, 'wb') as f:
            for chunk in file.chunks():
                f.write(chunk)

        research_agent = ResearchAgent(api_key=api_key)
        blog_agent = BlogAgent(api_key=api_key)

        content = research_agent.extract_audio_transcript(temp_path)
        os.remove(temp_path)

        blog_post = blog_agent.generate_blog(prompt, str(content))
        return {'content': blog_post}

    elif option == 'linkedin_post':
        if not file:
            return {'error': 'Audio file is required'}

        temp_path = f"temp_{file.name}"
        with open(temp_path, 'wb') as f:
            for chunk in file.chunks():
                f.write(chunk)

        research_agent = ResearchAgent(api_key=api_key)
        linkedin_agent = LinkedInAgent(api_key=api_key)

        content = research_agent.extract_audio_transcript(temp_path)
        os.remove(temp_path)

        linkedin_post = linkedin_agent.generate_linkedin_post(prompt, str(content))

        return {'content': linkedin_post}


# ______________________________________________________Ends Here and next api starts_________________________________
# -----------------------------------------------------------------------------------------------------------

# Dynamic agen creation code

# Create Dynamic Agent
@api_view(['POST'])
def create_dynamic_agent(request):
    """
    API to create an agent dynamically based on the dynamic_ai_agents table.
    """
    try:
        # Extract agent details from the request
        data = request.data
        name = data.get('agent_name')
        agent_goal = data.get('agent_goal')  # Change from 'agent_role' to 'agent_goal'
        agent_description = data.get('agent_description')
        agent_instruction = data.get('agent_instruction')
        email = data.get('email')  # Added email field

        # Ensure required fields are provided
        if not name or not agent_goal or not agent_description or not email:
            return Response({"error": "Agent name, goal, description, and email are required."}, status=400)

        # Step 1: Create the agent in the database
        agent_id = db.create_dynamic_agent(name, agent_goal, agent_description, agent_instruction, email)
        if not agent_id:
            return Response({"error": "Failed to create agent in the database."}, status=500)

        # Retrieve the created agent details
        agent = db.read_dynamic_agent(agent_id)
        if not agent:
            return Response({"error": "Agent details not found or incomplete."}, status=500)

        # Step 2: Construct agent details for the response
        agent_details = {
            "agent_id": agent_id,
            "agent_name": agent[1],
            "agent_goal": agent[2],
            "agent_description": agent[3],
            "agent_instruction": agent[4],
            "email": agent[5]  # Include email in the response
        }

        # Step 3: Return success response with agent details
        return Response({
            "message": "Agent created successfully.",
            "agent_details": agent_details
        }, status=201)

    except Exception as e:
        logger.error(f"Error creating agent: {str(e)}")  # Add detailed logging
        return Response({"error": f"Internal error: {str(e)}"}, status=400)


# Read Dynamic Agent by ID
@api_view(['GET'])
def read_dynamic_agent(request, agent_id):
    try:
        agent = db.read_dynamic_agent(agent_id)
        if agent:
            return Response({
                "id": agent[0],
                "agent_name": agent[1],
                "agent_goal": agent[2],
                "agent_description": agent[3],
                "agent_instruction": agent[4],
                "email": agent[5]  # Include email in the response
            }, status=200)
        return Response({"error": "Agent not found"}, status=404)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


# Update Dynamic Agent by ID
@api_view(['POST'])
def update_dynamic_agent(request, agent_id):
    try:
        data = request.data
        name = data.get('agent_name')
        agent_goal = data.get('agent_goal')
        agent_description = data.get('agent_description')
        agent_instruction = data.get('agent_instruction')
        email = data.get('email')  # Added email field

        # Update agent in the database
        db.update_dynamic_agent(agent_id, name, agent_goal, agent_description, agent_instruction, email)

        return Response({"message": f"Agent with ID {agent_id} updated successfully."}, status=200)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


# Delete Dynamic Agent by ID
@api_view(['GET'])
def delete_dynamic_agent(request, agent_id):
    try:
        db.delete_dynamic_agent(agent_id)
        return Response({"message": f"Agent with ID {agent_id} deleted successfully."}, status=204)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


@api_view(['POST'])
def get_all_dynamic_agents_by_email(request):
    """
    Fetch all dynamic agents associated with the given email.
    """
    try:
        email = request.POST.get('email')  # Extract email from query parameters
        if not email:
            return Response({"error": "Email is required."}, status=400)

        # Fetch dynamic agents from the database
        dynamic_agents = db.get_dynamic_agents_by_email(email)

        # Check if dynamic agents are found
        if not dynamic_agents:
            return Response({"message": "No dynamic agents found."}, status=404)

        # Structure the dynamic agents' data for JSON response
        dynamic_agents_data = [
            {
                "id": d_agent[0],
                "agent_name": d_agent[1],
                "agent_goal": d_agent[2],
                "agent_description": d_agent[3],
                "agent_instruction": d_agent[4],
                "email": d_agent[5]  # Include email in the response
            }
            for d_agent in dynamic_agents
        ]

        return Response({"dynamic_agents": dynamic_agents_data}, status=200)

    except Exception as e:
        # Log the error for further investigation
        logger.error(f"Error fetching dynamic agents: {str(e)}")
        return Response({"error": "An error occurred while fetching dynamic agents."}, status=500)


# Read All Dynamic Agents
@api_view(['GET'])
def read_all_dynamic_agents(request):
    try:
        # Fetch all agents from the database
        agents = db.get_all_dynamic_agents()

        # Check if any agents are returned
        if not agents:
            return Response({"message": "No agents found"}, status=404)

        # Structure the agents' data for JSON response
        agents_data = [
            {
                "id": agent[0],
                "agent_name": agent[1],
                "agent_goal": agent[2],
                "agent_description": agent[3],
                "agent_instruction": agent[4],
                "email": agent[5]  # Include email in the response
            }
            for agent in agents
        ]

        return Response({"agents": agents_data}, status=200)

    except Exception as e:
        # Log the error for further investigation
        logger.error(f"Error fetching agents: {str(e)}")

        # Return a user-friendly error message
        return Response({"error": "An error occurred while fetching agents"}, status=500)


# Main API to create an agent(dynamically)
@api_view(['POST'])
def create_openai_environment_api(request):
    """
    API to create an OpenAI environment dynamically based on agent_name, agent_goal, and agent_description.
    """
    try:
        agent_id = request.data.get('agent_id')

        if not agent_id:
            return Response({"error": "Agent ID is required"}, status=status.HTTP_400_BAD_REQUEST)

        # Retrieve the agent details from the database
        agent = db.read_dynamic_agent(agent_id)
        if not agent:
            return Response({"error": "Agent not found"}, status=status.HTTP_404_NOT_FOUND)

        # Extract agent details
        agent_details = {
            "name": agent[1],
            "agent_goal": agent[2],
            "agent_description": agent[3],
            "ext_tools": agent[4],
            "env_id": agent[5]
        }

        # Retrieve API key from the environment table
        env_details = db.read_environment(agent_details['env_id'])
        if not env_details or not env_details[2]:
            return Response({"error": "API key not found in environment table"},
                            status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        openai_api_key = env_details[2]

        # Dynamically construct the system prompt based on agent details
        system_prompt = generate_system_prompt(agent_details)

        # Add the system prompt to agent details for payload
        agent_details['system_prompt'] = system_prompt

        # Create the OpenAI environment
        environment_response = create_openai_environment(agent_details, openai_api_key)

        if environment_response.get("success"):
            return Response({"message": "OpenAI environment created successfully.",
                             "details": environment_response},
                            status=status.HTTP_201_CREATED)
        else:
            return Response({"error": "Failed to create OpenAI environment.",
                             "details": environment_response},
                            status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    except Exception as e:
        return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)


def generate_system_prompt(agent_details):
    return f"""
    You are {agent_details['name']}, an AI agent designed to achieve the following goal:
    {agent_details['agent_goal']}.

    Your description: {agent_details['agent_description']}.

    You have to follow the instructions:{agent_details['agent_instructions']}

    Behave in a manner consistent with this goal and description. Use external tools if required 
    (e.g., {', '.join(agent_details['tools']) if agent_details['tools'] else 'none'}). 
    Provide accurate, helpful, and concise responses to user queries.
    """


def create_openai_environment(agent_details, openai_api_key):
    try:
        url = "https://api.openai.com/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {openai_api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "gpt-4o-mini",  # Adjust the model based on your requirements
            "messages": [
                {"role": "system", "content": agent_details['system_prompt']},
                {"role": "user", "content": agent_details['agent_description']}
            ],
            "max_tokens": 1500
        }

        response = requests.post(url, headers=headers, json=payload)

        if response.status_code == 200:
            return {"success": True, "response": response.json()}
        else:
            return {"success": False, "error": response.text}

    except Exception as e:
        return {"success": False, "error": str(e)}


# from wyge.agents.prebuilt_agents import GoogleDriveAgent, EmailAgent
# from wyge.agents.prebuilt_agents import ResearchAgent, BlogAgent, LinkedInAgent, ImageGenerationAgent, EmailAgent


@api_view(['POST'])
def run_agent_environment(request):
    """
    API to interact with a dynamically created OpenAI environment using user-specific queries.
    """
    try:
        print("Received request to run_agent_environment")

        # Extract the agent ID and user query from the request
        data = request.data
        print("Request data:", data)

        agent_id = data.get('agent_id')
        user_query = data.get('query')
        recipient_email = data.get('email')  # Email comes from the frontend
        authorisation_code = data.get('auth_code')
        linkedin_token = data.get('token')

        if not agent_id or not user_query:
            print("Missing agent_id or query in request")
            return Response({"error": "Agent ID and query are required"}, status=400)

        print("Fetching agent details for agent_id:", agent_id)
        # Retrieve agent details from the database
        agent = db.read_agent(agent_id)
        if not agent:
            print("Agent not found for agent_id:", agent_id)
            return Response({"error": "Agent not found"}, status=404)

        dyn_agent = agent[8]
        dynamic_agent = db.read_dynamic_agent(dyn_agent)

        # Extract agent details
        agent_details = {
            "name": dynamic_agent[1],
            "agent_goal": dynamic_agent[2],
            "agent_description": dynamic_agent[3],
            "agent_instructions": dynamic_agent[4],
            "tools": agent[5]
        }
        print("Agent details:", agent_details)

        # Retrieve API key from the environment table
        print("Fetching environment details for env_id:", agent[7])
        env_details = db.read_environment(agent[7])
        if not env_details or not env_details[2]:
            print("API key not found in environment table for env_id:", agent_details['env_id'])
            return Response({"error": "API key not found in environment table"}, status=500)

        openai_api_key = env_details[2]
        print("OpenAI API key retrieved successfully")

        # Prepare payload for OpenAI API request
        url = "https://api.openai.com/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {openai_api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "gpt-4o-mini",  # Adjust the model based on your requirements
            "messages": [
                {"role": "system", "content": generate_system_prompt(agent_details)},
                {"role": "user", "content": user_query}
            ],
            "max_tokens": 1500
        }
        print("Payload prepared for OpenAI API request:", payload)

        # Send the query to the OpenAI environment
        print("Sending request to OpenAI API...")
        response = requests.post(url, headers=headers, json=payload)
        print("OpenAI API response status code:", response.status_code)

        if response.status_code == 200:
            openai_response = response.json()
            print("OpenAI API response:", openai_response)

            # Extract the 'choices' key from the OpenAI response to get the content
            chat_content = openai_response.get("choices", [{}])[0].get("message", {}).get("content", "")
            print("Extracted chat content:", chat_content)
            markdown_content = markdown_to_html(chat_content)

            # Check if any tools are specified in the agent
            if not agent[5]:  # No tools specified
                print("No tools specified. Returning OpenAI response directly.")
                return Response({
                    "success": True,
                    "message": "Response processed successfully.",
                    "content": markdown_content
                }, status=200)

            # Create a .docx file for the response
            file_name = "response.docx"
            file_path = f"./{file_name}"
            try:
                from docx import Document
                document = Document()
                document.add_heading("Response Content", level=1)
                document.add_paragraph(chat_content)
                document.save(file_path)
                print(f"{file_name} created successfully.")
            except Exception as e:
                return Response({"error": f"Failed to create .docx file: {str(e)}"}, status=500)

            # # Handle tools if specified
            # if "send_mail" in agent[5]:
            #     print("Using EmailAgent to send email...")
            #     email_agent = EmailAgent(api_key=openai_api_key)
            #     email_ack = email_agent.send_email(
            #         to_mail=recipient_email,
            #         subject=f"Response from {agent_details['name']}",
            #         body="Please find the response attached.",
            #         attachments=[file_path],
            #         credentials_json_file_path="credentials.json",
            #         token_json_file_path="token.json"
            #     )
            #     print("Email sent acknowledgment:", email_ack)
            #
            # elif "send_to_drive" in agent[5]:
            #     print("Handling Google Drive upload using user's authorization code...")
            #     drive_file_id = upload_to_drive(file_path, authorisation_code)
            #     if not drive_file_id:
            #         return Response({"error": "Failed to upload file to Google Drive"}, status=500)
            #
            # elif "linkedin_post" in agent[5]:
            #     print("Handling LinkedIn post...")
            #     linkedin_agent = LinkedInAgent(api_key=openai_api_key)
            #     linkedin_ack = linkedin_agent.post_content_on_linkedin(linkedin_token, markdown_content)
            #     print("LinkedIn post acknowledgment:", linkedin_ack)
            #
            # elif "google_calendar" in agent[5]:
            #     print("Handling Google Calendar scheduling...")
            #     event_details = {
            #         'summary': f"Event from {agent_details['name']}",
            #         'location': "Online",
            #         'description': chat_content,
            #         'time_zone': 'Asia/Kolkata',
            #         'start_time': datetime.utcnow() + timedelta(hours=1),
            #         'end_time': datetime.utcnow() + timedelta(hours=3)
            #     }
            #     calendar_event_link = create_google_calendar_event(event_details, authorisation_code)
            #     if not calendar_event_link:
            #         return Response({"error": "Failed to create Google Calendar event"}, status=500)
            #     print(f"Google Calendar event created: {calendar_event_link}")

            if "github" in agent[5]:
                print("Handling GitHub repository access...")
                repositories = list_github_repositories(authorisation_code)
                if not repositories:
                    return Response({"error": "Failed to fetch GitHub repositories"}, status=500)

            # Cleanup the local .docx file
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"Local file {file_name} deleted successfully.")

            return Response({
                "success": True,
                "message": "Response processed successfully.",
                "content": markdown_content
            }, status=200)
        else:
            print("Error response from OpenAI API:", response.text)
            return Response({
                "success": False,
                "error": response.text
            }, status=500)

    except Exception as e:
        print("Exception occurred:", str(e))
        return Response({"error": str(e)}, status=400)


import markdown


def markdown_to_html(md_text):
    html_text = markdown.markdown(md_text)
    return html_text


def upload_to_drive(file_path, authorization_code):
    try:
        import os
        from google_auth_oauthlib.flow import Flow
        from googleapiclient.discovery import build
        from googleapiclient.http import MediaFileUpload

        CLIENT_SECRETS_FILE = "google_credentials.json"
        REDIRECT_URI = "http://localhost:8000/oauth2callback/"

        # Debug: Print the file path and authorization code
        print(f"File path: {file_path}")
        print(f"Authorization code: {authorization_code}")

        if not os.path.exists(CLIENT_SECRETS_FILE):
            raise FileNotFoundError(f"Client secrets file not found at {CLIENT_SECRETS_FILE}")

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File to upload not found at {file_path}")

        # Exchange the authorization code for access tokens
        flow = Flow.from_client_secrets_file(
            CLIENT_SECRETS_FILE,
            scopes=["https://www.googleapis.com/auth/drive.file"],
            redirect_uri=REDIRECT_URI
        )
        flow.fetch_token(code=authorization_code)

        # Debug: Print the credentials information
        credentials = flow.credentials
        print(f"Access token: {credentials.token}")
        print(f"Refresh token: {credentials.refresh_token}")

        drive_service = build('drive', 'v3', credentials=credentials)

        # Debug: Print file metadata being used for upload
        file_metadata = {
            'name': os.path.basename(file_path),
        }
        print(f"File metadata: {file_metadata}")

        media = MediaFileUpload(
            file_path,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        print(f"Media upload object created: {media}")

        # Attempt to upload the file
        drive_file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id'
        ).execute()

        # Debug: Print the response from the API
        print(f"Drive API response: {drive_file}")

        print(f"File uploaded successfully to Google Drive with ID: {drive_file.get('id')}")
        return drive_file.get('id')

    except Exception as e:
        print(f"Failed to upload file to Google Drive: {str(e)}")
        return None


def create_google_calendar_event(event_details, authorization_code):
    try:
        import os
        from google_auth_oauthlib.flow import Flow
        from googleapiclient.discovery import build

        CLIENT_SECRETS_FILE = "google_credentials.json"
        REDIRECT_URI = "http://localhost:8000/oauth2callback/"
        SCOPES = ['https://www.googleapis.com/auth/calendar']

        # Debug: Print event details and authorization code
        print(f"Event details: {event_details}")
        print(f"Authorization code: {authorization_code}")

        if not os.path.exists(CLIENT_SECRETS_FILE):
            raise FileNotFoundError(f"Client secrets file not found at {CLIENT_SECRETS_FILE}")

        # Exchange the authorization code for access tokens
        flow = Flow.from_client_secrets_file(
            CLIENT_SECRETS_FILE,
            scopes=SCOPES,
            redirect_uri=REDIRECT_URI
        )
        flow.fetch_token(code=authorization_code)

        # Debug: Print the credentials information
        credentials = flow.credentials
        print(f"Access token: {credentials.token}")
        print(f"Refresh token: {credentials.refresh_token}")

        # Build the Google Calendar service
        service = build('calendar', 'v3', credentials=credentials)

        # Ensure that start and end time are included in the event details
        event_start_time = event_details.get('start_time', datetime.utcnow())
        event_end_time = event_details.get('end_time', event_start_time + timedelta(hours=1))  # Default to 1-hour event

        # Format the times in ISO 8601 format
        start_time_iso = event_start_time.isoformat() + 'Z'
        end_time_iso = event_end_time.isoformat() + 'Z'

        # Construct the event body
        event = {
            'summary': event_details.get('summary', 'No Title'),
            'location': event_details.get('location', ''),
            'description': event_details.get('description', ''),
            'start': {'dateTime': start_time_iso, 'timeZone': 'UTC'},
            'end': {'dateTime': end_time_iso, 'timeZone': 'UTC'},
            # 'attendees': [{'email': email} for email in event_details.get('attendees', [])],
            'reminders': {
                'useDefault': False,
                'overrides': [
                    {'method': 'email', 'minutes': 24 * 60},
                    {'method': 'popup', 'minutes': 10},
                ],
            },
        }

        # Create the event
        created_event = service.events().insert(calendarId='primary', body=event).execute()

        # Debug: Print the created event response
        print(f"Event created successfully: {created_event['htmlLink']}")
        return created_event['htmlLink']

    except Exception as e:
        print(f"Failed to create Google Calendar event: {str(e)}")
        return None


def list_github_repositories(authorization_code):
    try:
        import requests

        CLIENT_ID = os.getenv("CLIENT_ID")  # Replace with your GitHub OAuth app client ID
        CLIENT_SECRET = os.getenv("CLIENT_SECRET")  # Replace with your GitHub OAuth app client secret
        REDIRECT_URI = "http://localhost:8000/oauth2callback/"
        GITHUB_TOKEN_URL = "https://github.com/login/oauth/access_token"
        GITHUB_API_URL = "https://api.github.com"

        # Debug: Print authorization code
        print(f"Authorization code: {authorization_code}")

        # Exchange the authorization code for an access token
        token_response = requests.post(
            GITHUB_TOKEN_URL,
            headers={"Accept": "application/json"},
            data={
                "client_id": CLIENT_ID,
                "client_secret": CLIENT_SECRET,
                "code": authorization_code,
                "redirect_uri": REDIRECT_URI,
            },
        )
        if token_response.status_code != 200:
            raise Exception(f"Failed to fetch access token: {token_response.json()}")

        token_data = token_response.json()
        access_token = token_data.get("access_token")

        # Debug: Print access token
        print(f"Access token: {access_token}")

        if not access_token:
            raise Exception("Access token is missing in the response.")

        # Use the access token to make authenticated requests to the GitHub API
        headers = {"Authorization": f"Bearer {access_token}"}

        # Fetch the list of repositories for the authenticated user
        response = requests.get(f"{GITHUB_API_URL}/user/repos", headers=headers)

        if response.status_code != 200:
            raise Exception(f"Failed to fetch repositories: {response.json()}")

        repos = response.json()

        # Debug: Print the list of repositories
        print("Repositories fetched successfully:")
        for repo in repos:
            print(f"- {repo['name']} (URL: {repo['html_url']})")

        # Return the list of repository names and URLs
        return [{"name": repo["name"], "url": repo["html_url"]} for repo in repos]

    except Exception as e:
        print(f"Failed to fetch GitHub repositories: {str(e)}")
        return None
