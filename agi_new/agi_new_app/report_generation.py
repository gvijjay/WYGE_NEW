# # Environment Creation
# import os
#
# import numpy as np
# from django.core.files.storage import default_storage
# from django.http import JsonResponse
# from django.views.decorators.csrf import csrf_exempt
# from openai import OpenAI
# from plotly.graph_objs import Figure
# from rest_framework import status
# from rest_framework.response import Response
# from rest_framework.decorators import api_view
#
# from .database import PostgreSQLDB
# from ..agi_new import settings
#
# db = PostgreSQLDB(dbname='test', user='test_owner', password='tcWI7unQ6REA')
#
# @api_view(['POST'])
# def run_openai_environment_report(request):
#     try:
#         agent_id = request.data.get('agent_id')
#         user_prompt = request.data.get('prompt', '')
#         file = request.FILES.get('file')
#         url = request.data.get('url', '')
#         file1 = request.FILES.getlist("file")
#
#         # Retrieve agent details
#         agent = db.read_agent(agent_id)
#
#         # Retrieve the API key from the environment table using env_id
#         env_details = db.read_environment(agent[7])
#         openai_api_key = env_details[2]
#         client = OpenAI(api_key=openai_api_key)
#
#         # First Application: Text_to_sql
#         if file:
#             print("Report_Generation_condition")
#             result = gen_response(file, user_prompt, client)
#             # Handle the result from gen_response
#             if not result:
#                 return Response({"error": "gen_response() returned None."},
#                                 status=status.HTTP_500_INTERNAL_SERVER_ERROR)
#             if "chartData" in result:
#                 # Visualization result
#                 return Response({"chartData": result["chartData"]}, status=status.HTTP_200_OK)
#             elif "answer" in result:
#                 # Text-based result
#                 return Response({"answer": markdown_to_html(result["answer"])}, status=status.HTTP_200_OK)
#             else:
#                 return Response({"error": "No valid output from gen_response."}, status=status.HTTP_400_BAD_REQUEST)
#         else:
#             return Response({"error": "No valid tool found for the given input."}, status=status.HTTP_400_BAD_REQUEST)
#
#     except Exception as e:
#         return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
#
#
# import markdown
# def markdown_to_html(md_text):
#     html_text = markdown.markdown(md_text)
#     return html_text
#
#
# ##1st Application:Text-to-sql(Both text and graph)
# USER = 'test_owner'
# PASSWORD = 'tcWI7unQ6REA'
# HOST = 'ep-yellow-recipe-a5fny139.us-east-2.aws.neon.tech:5432'
# DATABASE = 'test'
#
# @csrf_exempt
# def gen_response(file, query, client):
#     table_name = file.name.split('.')[0]
#     print(table_name)
#
#     # Define upload directory
#     upload_dir = os.path.join(settings.MEDIA_ROOT, 'upload')
#     print(upload_dir)
#     os.makedirs(upload_dir, exist_ok=True)
#
#     # File path within upload directory
#     file_path = os.path.join(upload_dir, file.name)
#     with default_storage.open(file_path, 'wb+') as f:
#         for chunk in file.chunks():
#             f.write(chunk)
#
#     df = file_to_sql(file_path, table_name, USER, PASSWORD, HOST, DATABASE)
#     print(df.head(5))
#
#     # Generate CSV metadata
#     csv_metadata = {"columns": df.columns.tolist()}
#     metadata_str = ", ".join(csv_metadata["columns"])
#     if not query:
#         return JsonResponse({"error": "No query provided"}, status=400)
#
#     graph_keywords = [
#         "plot", "graph", "visualize", "visualization", "scatter", "bar chart",
#         "line chart", "histogram", "pie chart", "bubble chart", "heatmap", "box plot",
#         "generate chart", "create graph", "draw", "trend", "correlation"
#     ]
#
#     # Decide whether the query is text-based or graph-based
#     if any(keyword in query.lower() for keyword in graph_keywords):
#         # Graph-related prompt
#         print("if_condition------------------")
#         print(query)
#         prompt_eng = (
#             f"You are an AI specialized in data analytics and visualization."
#             f"Data used for analysis is stored in a pandas DataFrame named `df`. "
#             f"The DataFrame `df` contains the following attributes: {metadata_str}. "
#             f"Based on the user's query, generate Python code using Plotly to create the requested type of graph "
#             f"(e.g., bar, pie, scatter, etc.) using the data in the DataFrame `df`. "
#             f"The graph must utilize the data within `df` as appropriate for the query. "
#             f"If the user does not specify a graph type, decide whether to generate a line or bar graph based on the situation."
#             f"Every graph must include a title, axis labels (if applicable), and appropriate colors for better visualization."
#             f"The graph must have a white background for both the plot and paper. "
#             f"The code must output a Plotly 'Figure' object stored in a variable named 'fig', "
#             f"and include 'data' and 'layout' dictionaries compatible with React. "
#             f"The user asks: {query}."
#         )
#
#         # Call AI to generate the code
#         chat = generate_code(prompt_eng, client)
#         print("Generated code from AI:")
#         print(chat)
#
#         # Check for valid Plotly code in the AI response
#         if 'import' in chat:
#             namespace = {'df': df}  # Pass `df` into the namespace
#             try:
#                 # Execute the generated code
#                 exec(chat, namespace)
#
#                 # Retrieve the Plotly figure from the namespace
#                 fig = namespace.get("fig")
#
#                 if fig and isinstance(fig, Figure):
#                     # Convert the Plotly figure to JSON
#                     chart_data = fig.to_plotly_json()
#
#                     # Ensure JSON serialization by converting NumPy arrays to lists
#                     def make_serializable(obj):
#                         if isinstance(obj, np.ndarray):
#                             return obj.tolist()
#                         elif isinstance(obj, dict):
#                             return {k: make_serializable(v) for k, v in obj.items()}
#                         elif isinstance(obj, list):
#                             return [make_serializable(v) for v in obj]
#                         return obj
#
#                     # Recursively process the chart_data
#                     chart_data_serializable = make_serializable(chart_data)
#
#                     # Return the structured response to the frontend
#                     return {"chartData": chart_data_serializable}
#                 else:
#                     return {"message": "No valid Plotly figure found."}
#             except Exception as e:
#                 return {"message": f"Error executing code: {str(e)}"}
#         else:
#             return {"message": "AI response does not contain valid code."}
#
#
#
# # Function to generate code from OpenAI API
# def generate_code(prompt_eng, client):
#     response = client.chat.completions.create(
#         model="gpt-4o-mini",
#         messages=[
#             {"role": "system", "content": "You are a helpful assistant."},
#             {"role": "user", "content": prompt_eng}
#         ]
#     )
#     all_text = ""
#     for choice in response.choices:
#         message = choice.message
#         chunk_message = message.content if message else ''
#         all_text += chunk_message
#     print(all_text)
#     if "```python" in all_text:
#         code_start = all_text.find("```python") + 9
#         code_end = all_text.find("```", code_start)
#         code = all_text[code_start:code_end]
#     else:
#         code = all_text
#     return code
#
#
#
# def file_to_sql(file_path, table_name, user, password, host, db_name):
#     import pandas as pd
#     import os
#     from sqlalchemy import create_engine
#
#     # engine = create_engine(f"postgresql://{user}:{password}@{host}/{db_name}")
#     engine = create_mysql_engine(user, password, host, db_name)
#
#     if not table_name:
#         table_name = os.path.splitext(os.path.basename(file_path))[0]
#
#     file_extension = os.path.splitext(file_path)[-1].lower()
#     if file_extension == '.xlsx':
#         df = pd.read_excel(file_path)
#     elif file_extension == '.csv':
#         df = pd.read_csv(file_path)
#     else:
#         raise ValueError("Unsupported file format. Please provide an Excel (.xlsx) or CSV (.csv) file.")
#
#     df.to_sql(table_name, con=engine, if_exists='replace', index=False)
#     return df
#
#
# def create_mysql_engine(user, password, host, db_name):
#     from sqlalchemy import create_engine, text
#
#     if db_name:
#         connection_str = f'postgresql://{user}:{password}@{host}/{db_name}'
#     else:
#         connection_str = f'postgresql://{user}:{password}@{host}/'
#     engine = create_engine(connection_str)
#     return engine

# Report Generation code
import os
from datetime import datetime

import pandas as pd
import numpy as np
import markdown
from django.conf import settings
from django.core.files.storage import default_storage
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from openai import OpenAI
from plotly.graph_objs import Figure
from rest_framework import status
from rest_framework.response import Response
from rest_framework.decorators import api_view
from docx import Document
from docx.shared import Inches
import plotly.io as pio
import tempfile
from sqlalchemy import create_engine

from .database import PostgreSQLDB

db = PostgreSQLDB(dbname='test', user='test_owner', password='tcWI7unQ6REA')


@api_view(['POST'])
def run_openai_environment_report(request):
    try:
        agent_id = request.data.get('agent_id')
        user_prompt = request.data.get('prompt', 'Generate a comprehensive report with visualizations')
        file = request.FILES.get('file')

        if not file:
            return Response({"error": "No file provided"}, status=status.HTTP_400_BAD_REQUEST)

        # Retrieve agent details
        agent = db.read_agent(agent_id)
        if not agent:
            return Response({"error": "Agent not found"}, status=status.HTTP_404_NOT_FOUND)

        # Retrieve the API key from the environment table using env_id
        env_details = db.read_environment(agent[7])
        if not env_details:
            return Response({"error": "Environment not found"}, status=status.HTTP_404_NOT_FOUND)

        openai_api_key = env_details[2]
        client = OpenAI(api_key=openai_api_key)

        # Process the file and generate report
        report_data = generate_report_with_graphs(file, user_prompt, client)

        if not report_data or not report_data.get('success'):
            error_msg = report_data.get('error',
                                        'Failed to generate report') if report_data else 'Failed to generate report'
            return Response({"error": error_msg},
                            status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # Prepare the response data
        response_data = {
            "message": "Report generated successfully",
            "content_summary":markdown_to_html( report_data['content']),
            "visualizations": report_data['visualization_descriptions']
        }

        # If you want to include the actual file content in the response
        if settings.DEBUG:  # Only in development for testing
            try:
                with open(report_data['file_path'], 'rb') as f:
                    response_data['file_content'] = f.read().decode('latin-1')
            except Exception as e:
                response_data['file_read_error'] = str(e)

        return Response(response_data, status=status.HTTP_200_OK)

    except Exception as e:
        return Response({
            "error": str(e),
            "details": "An error occurred during report generation"
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

def markdown_to_html(md_text):
    html_text = markdown.markdown(md_text)
    return html_text


def generate_report_with_graphs(file, user_prompt, client):
    """Process Excel file, generate report with graphs, and return serializable data"""
    result = {
        'content': None,
        'file_path': None,
        'visualization_descriptions': [],
        'success': False,
        'error': None
    }

    try:
        # Step 1: Load the Excel file
        print("Excel_file Loading..")
        df = load_excel_file(file)
        print("Excel_file_loaded")

        if df.empty:
            raise ValueError("Empty DataFrame after loading Excel file")

        # Step 2: Generate analysis text
        print("Analysis going to happen")
        analysis_text = generate_analysis_text(df, user_prompt, client)
        print("Analysis has been done")

        # Step 3: Generate visualizations
        print("Visualisation Going to happen")
        visualizations = generate_visualizations(df, user_prompt, client)
        print(f"Visualisation done ({len(visualizations)} graphs generated)")

        # Store visualization descriptions instead of figure objects
        for idx, fig in enumerate(visualizations):
            description = {
                'figure_number': idx + 1,
                'title': None,
                'description': None
            }

            if hasattr(fig, '_suptitle') and fig._suptitle is not None:
                description['title'] = fig._suptitle.get_text()
            elif len(fig.axes) > 0 and fig.axes[0].get_title():
                description['title'] = fig.axes[0].get_title()

            # Add simple description based on visualization type
            if len(fig.axes) > 0:
                for artist in fig.axes[0].get_children():
                    if hasattr(artist, 'get_label'):
                        description['description'] = f"Contains {artist.__class__.__name__} visualization"
                        break

            result['visualization_descriptions'].append(description)

        # Step 4: Create Word document
        print("Document Creation")
        doc = Document()

        # Add title and metadata
        doc.add_heading('Analysis Report', level=0)

        # Add analysis sections
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(analysis_text['executive_summary'])

        doc.add_heading('Key Findings', level=1)
        doc.add_paragraph(analysis_text['key_findings'])

        doc.add_heading('Detailed Analysis', level=1)
        doc.add_paragraph(analysis_text['detailed_analysis'])

        # Add visualizations if any were generated
        if visualizations:
            doc.add_heading('Data Visualizations', level=1)
            for idx, fig in enumerate(visualizations):
                try:
                    # Save plot to temporary image file
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                        fig.savefig(tmp.name, format='png', dpi=300, bbox_inches='tight')
                        tmp_path = tmp.name

                    # Add image to document
                    doc.add_picture(tmp_path, width=Inches(4.75))

                    # Add figure caption
                    fig_title = f"Figure {idx + 1}"
                    if hasattr(fig, '_suptitle') and fig._suptitle is not None:
                        fig_title += f": {fig._suptitle.get_text()}"
                    elif len(fig.axes) > 0 and fig.axes[0].get_title():
                        fig_title += f": {fig.axes[0].get_title()}"

                    doc.add_paragraph(fig_title)
                    doc.add_paragraph("\n")

                    # Clean up temporary file
                    os.unlink(tmp_path)
                except Exception as e:
                    print(f"Error adding visualization {idx}: {str(e)}")
                    continue

        # Add recommendations if available
        if 'recommendations' in analysis_text and analysis_text['recommendations']:
            doc.add_heading('Recommendations', level=1)
            doc.add_paragraph(analysis_text['recommendations'])

        # Save the document
        report_dir = os.path.join(settings.MEDIA_ROOT, 'reports')
        os.makedirs(report_dir, exist_ok=True)
        report_filename = f"report_{file.name.split('.')[0]}.docx"
        report_path = os.path.join(report_dir, report_filename)
        doc.save(report_path)
        result['file_path'] = report_path

        # Extract text content from the document (serializable)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        result['content'] = "\n".join(full_text)
        result['success'] = True

        print(f"Report successfully generated at: {report_path}")
        return result

    except Exception as e:
        error_msg = f"Error in generate_report_with_graphs: {str(e)}"
        print(error_msg)
        result['error'] = error_msg
        return result

def load_excel_file(file):
    """Load Excel file into DataFrame"""
    try:
        # Save uploaded file temporarily
        upload_dir = os.path.join(settings.MEDIA_ROOT, 'temp_uploads')
        os.makedirs(upload_dir, exist_ok=True)
        file_path = os.path.join(upload_dir, file.name)

        with default_storage.open(file_path, 'wb+') as f:
            for chunk in file.chunks():
                f.write(chunk)

        # Read Excel file
        df = pd.read_excel(file_path)

        # Clean up temporary file
        os.unlink(file_path)

        return df

    except Exception as e:
        print(f"Error loading Excel file: {str(e)}")
        return pd.DataFrame()


def generate_analysis_text(df, user_prompt, client):
    """Generate analysis text using LLM"""
    try:
        # Get basic statistics about the DataFrame
        stats = df.describe().to_string()
        columns = ", ".join(df.columns.tolist())

        prompt = f"""
        You are a data analysis expert. Analyze the following dataset and generate a comprehensive report.

        Dataset Information:
        - Columns: {columns}
        - Shape: {df.shape}
        - Basic Statistics:
        {stats}

        User's request: {user_prompt}

        Generate a report with the following sections:

        === EXECUTIVE SUMMARY ===
        [3-5 sentences summarizing key insights]

        === KEY FINDINGS ===
        [bullet points of most important findings]

        === DETAILED ANALYSIS ===
        [paragraphs explaining the data in detail]

        === RECOMMENDATIONS ===
        [actionable suggestions based on the data]

        Please format your response exactly as shown above with the section headers.
        """

        response = client.chat.completions.create(
            model="gpt-4o-mini",  # Use your actual model name
            messages=[
                {"role": "system", "content": "You are a helpful data analysis assistant."},
                {"role": "user", "content": prompt}
            ]
        )

        # Parse the response text
        response_text = response.choices[0].message.content

        # Parse the sections from the response
        sections = {
            "executive_summary": "",
            "key_findings": "",
            "detailed_analysis": "",
            "recommendations": ""
        }

        current_section = None
        for line in response_text.split('\n'):
            if line.startswith("=== EXECUTIVE SUMMARY ==="):
                current_section = "executive_summary"
            elif line.startswith("=== KEY FINDINGS ==="):
                current_section = "key_findings"
            elif line.startswith("=== DETAILED ANALYSIS ==="):
                current_section = "detailed_analysis"
            elif line.startswith("=== RECOMMENDATIONS ==="):
                current_section = "recommendations"
            elif current_section and line.strip() and not line.startswith("==="):
                sections[current_section] += line + "\n"

        return sections

    except Exception as e:
        print(f"Error generating analysis text: {str(e)}")
        return {
            "executive_summary": "Analysis could not be generated",
            "key_findings": "Analysis could not be generated",
            "detailed_analysis": "Analysis could not be generated",
            "recommendations": "Analysis could not be generated"
        }


import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO
from matplotlib.figure import Figure


def generate_visualizations(df, user_prompt, client, num_graphs=3):
    """Generate Matplotlib visualizations based on the data"""
    figures = []

    # Generate different visualizations
    for i in range(num_graphs):
        try:
            # Get a different prompt variation for each graph
            graph_prompt = get_graph_prompt(i, user_prompt, df)

            # Generate code for the visualization
            code = generate_matplotlib_code(df, graph_prompt, client)

            if not code:
                continue

            # Execute the code to get the figure
            namespace = {'df': df, 'plt': plt, 'np': np, 'BytesIO': BytesIO}
            exec(code, namespace)
            fig = namespace.get('fig')

            if fig and isinstance(fig, Figure):
                figures.append(fig)
        except Exception as e:
            print(f"Error generating visualization {i}: {str(e)}")
            continue

    return figures


def save_matplotlib_figure(fig):
    """Save matplotlib figure to bytes"""
    buf = BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight', dpi=100)
    buf.seek(0)
    return buf


def get_graph_prompt(index, base_prompt, df):
    """Get different prompt variations for different graphs"""
    columns = ", ".join(df.columns.tolist())
    num_cols = len(df.columns)

    # Dynamic visualization types based on data characteristics
    if len(df) > 1000:
        size_note = "Since we have substantial data points, consider using density plots or aggregated views."
    else:
        size_note = "With this amount of data, individual data points can be clearly visualized."

    # Different visualization styles for each index
    styles = [
        "modern business style with clean lines and corporate colors",
        "academic publication style with detailed annotations",
        "infographic style with bold colors and clear highlights"
    ]

    prompts = [
        f"""Create an advanced visualization showing complex relationships in the data. Columns: {columns}.
        {size_note}
        Suggestions: Small multiples, paired plots, or correlation matrix.
        Style: {styles[index % 3]}.
        Make it publication-quality with proper annotations.""",

        f"""Generate a sophisticated comparative visualization. Columns: {columns}.
        {size_note}
        Suggestions: Stacked area charts, radar charts, or violin plots.
        Style: {styles[(index + 1) % 3]}.
        Include statistical annotations if appropriate.""",

        f"""Create an innovative visualization that reveals hidden patterns. Columns: {columns}.
        {size_note}
        Suggestions: Hexbin plots, 2D histograms, or parallel coordinates.
        Style: {styles[(index + 2) % 3]}.
        Make it visually striking but informative."""
    ]

    # Add the user's specific request to each prompt
    enhanced_prompt = f"{base_prompt}. Focus on creating a unique, non-standard visualization that provides novel insights."

    if index < len(prompts):
        return f"{enhanced_prompt} {prompts[index]}"
    return enhanced_prompt


def generate_matplotlib_code(df, prompt, client):
    """Generate Matplotlib visualization code using LLM"""
    columns = ", ".join(df.columns.tolist())
    sample_data = df.head(3).to_string()

    full_prompt = f"""
    You are a data visualization expert. Create an advanced Matplotlib visualization that stands out.

    Dataset Information:
    - Columns: {columns}
    - Shape: {df.shape}
    - Sample data: {sample_data}
    - Data types: {df.dtypes.to_string()}

    User's request: {prompt}

    Requirements:
    1. Create a NON-DEFAULT, sophisticated visualization
    2. Begin with 'fig, ax = plt.subplots(figsize=(10,6))' or similar
    3. Use advanced features like:
       - Custom annotations
       - Statistical overlays
       - Multiple axes if appropriate
       - Professional color palettes
    4. Include:
       - Informative title (fig.suptitle())
       - Axis labels with units if applicable
       - Legend if needed
       - Grid or reference lines
    5. Style considerations:
       - Use plt.style.context('seaborn-v0_8') or similar
       - Customize tick params
       - Adjust spines
    6. Preprocessing:
       - Handle missing values if needed
       - Consider log scales if appropriate
    7. Return ONLY the code wrapped in ```python ``` markers.

    Example structure:
    ```python
    with plt.style.context('seaborn-v0_8'):
        fig, ax = plt.subplots(figsize=(10, 6))
        # Visualization code here
        ax.set_title(...)
        # Custom styling here
    ```
    """

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a creative visualization designer."},
            {"role": "user", "content": full_prompt}
        ]
    )


    code = response.choices[0].message.content

    if "```python" in code:
        code = code.split("```python")[1].split("```")[0]
    return code




